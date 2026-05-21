"""
生成Demo素材：
1. 主文档：网销渠道退保功能需求规格说明书.docx
2. 附件A：退保数据字典.xlsx
3. 附件B：退保接口规范.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "input")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(text)
        for p in cell.paragraphs:
            p.style.font.size = Pt(10)
            if bold:
                for run in p.runs:
                    run.bold = True
        if header:
            set_cell_shading(cell, "4472C4")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True


# ============================================================
# 主文档：网销渠道退保功能需求规格说明书
# ============================================================
def generate_main_docx():
    doc = Document()

    # --- 标题 ---
    title = doc.add_heading("网销渠道退保功能需求规格说明书", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 文档信息表
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "文档编号"
    hdr[1].text = "REQ-2026-05-001"
    hdr[2].text = "版本"
    hdr[3].text = "V1.0"
    row = table.add_row().cells
    row[0].text = "编写人"
    row[1].text = "产品部-张三"
    row[2].text = "编写日期"
    row[3].text = "2026-05-22"
    row = table.add_row().cells
    row[0].text = "审核人"
    row[1].text = "技术部-李四"
    row[2].text = "状态"
    row[3].text = "评审中"

    doc.add_paragraph("")

    # === 1. 需求背景 ===
    doc.add_heading("1. 需求背景", level=1)
    doc.add_heading("1.1 业务现状", level=2)
    doc.add_paragraph(
        "目前公司网销渠道（官方APP、微信小程序）的退保业务仍依赖线下柜面办理，"
        "投保人需携带身份证件前往柜台填写退保申请书，由柜员手工录入系统后发起退保流程。"
        "该模式存在以下问题："
    )
    doc.add_paragraph("客户体验差：投保人须在工作时间前往柜面，排队等候时间长", style="List Bullet")
    doc.add_paragraph("运营成本高：每笔退保需柜员人工处理，人力成本持续上升", style="List Bullet")
    doc.add_paragraph("时效不可控：柜面受理后还需流转至保全岗审核，全流程平均耗时5个工作日", style="List Bullet")

    doc.add_heading("1.2 需求目标", level=2)
    doc.add_paragraph(
        "实现网销渠道退保全流程线上化，投保人通过APP或小程序即可自助完成退保申请，"
        "系统自动完成身份校验、退费金额计算和退费执行，将退保处理时效从5个工作日缩短至实时处理。"
    )

    # === 2. 业务流程 ===
    doc.add_heading("2. 业务流程", level=1)
    doc.add_heading("2.1 整体流程", level=2)
    doc.add_paragraph("退保业务的整体流程如下图所示：")

    # Mermaid 流程图（文本形式）
    mermaid = """
```mermaid
flowchart TD
    A[投保人发起退保申请] --> B{保单状态校验}
    B -->|有效| C{是否在犹豫期内}
    B -->|中止/终止| D[提示不可退保]
    C -->|是| E[按犹豫期退保计算]
    C -->|否| F[按正常退保计算]
    E --> G[退费金额确认]
    F --> G
    G --> H[投保人确认退保]
    H --> I[调用收付微服务执行退费]
    I --> J[更新保单状态为终止]
    J --> K[发送退保确认通知]
```"""
    doc.add_paragraph(mermaid)

    doc.add_heading("2.2 犹豫期退保规则", level=2)
    doc.add_paragraph(
        "投保人签收电子保单后15天内为犹豫期。犹豫期内申请退保的，退费金额计算公式为："
    )
    doc.add_paragraph("退费金额 = 已缴保费 - 工本费（不超过10元）")
    doc.add_paragraph("犹豫期退保的退费须在受理后5个工作日内完成，退费至投保人原缴费账户。")

    doc.add_heading("2.3 正常退保规则", level=2)
    doc.add_paragraph(
        "超过犹豫期后的退保，退还保单的现金价值。现金价值根据保单已生效年限、"
        "已缴保费总额、险种类型等因素计算，具体金额参见产品条款中的现金价值表。"
        "业务人员须向投保人充分告知退保损失。"
    )
    doc.add_paragraph(
        "如保单存在保单贷款未还清的情况，退费金额须先扣除贷款本息。"
    )

    # === 3. 功能需求 ===
    doc.add_heading("3. 功能需求", level=1)
    doc.add_heading("3.1 退保申请", level=2)
    doc.add_paragraph("投保人通过APP或小程序提交退保申请，系统须完成以下校验：")

    # 校验规则表
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "序号"
    hdr[1].text = "校验项"
    hdr[2].text = "校验规则"
    hdr[3].text = "不通过处理"
    set_cell_shading(hdr[0], "4472C4")
    set_cell_shading(hdr[1], "4472C4")
    set_cell_shading(hdr[2], "4472C4")
    set_cell_shading(hdr[3], "4472C4")

    rules = [
        ("1", "保单状态", "保单必须为有效状态", "提示：保单当前状态不可退保"),
        ("2", "申请人身份", "申请人须为投保人本人", "提示：仅投保人可申请退保"),
        ("3", "犹豫期判断", "判断当前是否在犹豫期内", "按对应规则计算退费金额"),
        ("4", "保单贷款", "检查是否存在未还清的保单贷款", "退费金额先扣除贷款本息"),
        ("5", "受益人同意", "含身故受益人的保单须所有受益人同意", "提示：请联系受益人确认"),
        ("6", "反洗钱校验", "退保金额超过5万元触发身份重新识别", "引导投保人完成身份核验"),
    ]
    for r in rules:
        add_table_row(table, r)

    doc.add_heading("3.2 退费计算", level=2)
    doc.add_paragraph("系统须根据退保类型自动计算退费金额：")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "退保类型"
    hdr[1].text = "退费公式"
    hdr[2].text = "退费时效"
    set_cell_shading(hdr[0], "4472C4")
    set_cell_shading(hdr[1], "4472C4")
    set_cell_shading(hdr[2], "4472C4")

    calcs = [
        ("犹豫期退保", "已缴保费 - 工本费", "受理后5个工作日"),
        ("正常退保", "现金价值 - 保单贷款本息", "受理后10个工作日"),
    ]
    for c in calcs:
        add_table_row(table, c)

    doc.add_heading("3.3 通知机制", level=2)
    doc.add_paragraph("退保完成后系统须向投保人发送以下通知：")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "通知节点"
    hdr[1].text = "通知方式"
    hdr[2].text = "通知内容"
    set_cell_shading(hdr[0], "4472C4")
    set_cell_shading(hdr[1], "4472C4")
    set_cell_shading(hdr[2], "4472C4")

    notices = [
        ("退保申请受理", "APP推送+短信", "您的退保申请已受理，预计X个工作日完成退费"),
        ("退费执行成功", "APP推送+短信", "退费金额XX元已退还至您的XX账户"),
        ("退保完成", "APP推送+短信", "您的保单已退保终止，如需保障请重新投保"),
    ]
    for n in notices:
        add_table_row(table, n)

    # === 4. 非功能需求 ===
    doc.add_heading("4. 非功能需求", level=1)
    doc.add_heading("4.1 性能要求", level=2)
    doc.add_paragraph("退保申请接口响应时间不超过3秒", style="List Bullet")
    doc.add_paragraph("退费执行接口响应时间不超过5秒", style="List Bullet")
    doc.add_paragraph("系统须支持日均5000笔退保申请的并发处理", style="List Bullet")

    doc.add_heading("4.2 安全要求", level=2)
    doc.add_paragraph("退保操作须进行投保人身份二次验证（短信验证码或人脸识别）", style="List Bullet")
    doc.add_paragraph("退保金额超过5万元须触发反洗钱客户身份重新识别", style="List Bullet")
    doc.add_paragraph("退保操作日志须完整记录，保留期不少于10年", style="List Bullet")

    # === 5. 数据要求 ===
    doc.add_heading("5. 数据要求", level=1)
    doc.add_paragraph(
        "退保操作涉及的数据字段及存储结构详见附件A《退保数据字典》。"
        "退保相关的微服务接口定义详见附件B《退保接口规范》。"
    )

    # === 6. 验收标准 ===
    doc.add_heading("6. 验收标准", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "编号"
    hdr[1].text = "验收项"
    hdr[2].text = "验收标准"
    set_cell_shading(hdr[0], "4472C4")
    set_cell_shading(hdr[1], "4472C4")
    set_cell_shading(hdr[2], "4472C4")

    acpts = [
        ("AC-01", "犹豫期退保", "投保人在犹豫期内退保，退费金额等于已缴保费减工本费"),
        ("AC-02", "正常退保", "超过犹豫期退保，退费金额等于保单现金价值"),
        ("AC-03", "退费时效", "退费在受理后规定工作日内到达投保人账户"),
        ("AC-04", "通知送达", "退保各节点通知通过APP推送和短信送达投保人"),
        ("AC-05", "保单状态", "退保完成后保单状态变更为终止"),
        ("AC-06", "异常处理", "退费失败时系统自动重试并通知运维人员"),
    ]
    for a in acpts:
        add_table_row(table, a)

    path = os.path.join(OUTPUT_DIR, "网销渠道退保功能需求规格说明书.docx")
    doc.save(path)
    print(f"[OK] 主文档已生成: {path}")


# ============================================================
# 附件A：退保数据字典.xlsx
# ============================================================
def generate_data_dict_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "退保数据字典"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    # Sheet1: 保单主表字段
    headers = ["字段名", "字段类型", "长度", "必填", "说明", "示例值"]
    ws.append(headers)
    for col_idx, _ in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    data = [
        ("policy_no", "VARCHAR", "32", "Y", "保单号，主键", "PL2026050001"),
        ("policy_holder_id", "VARCHAR", "32", "Y", "投保人ID", "H20260001"),
        ("insured_id", "VARCHAR", "32", "Y", "被保险人ID", "I20260001"),
        ("status", "VARCHAR", "16", "Y", "保单状态：VALID/SUSPENDED/TERMINATED", "VALID"),
        ("effective_date", "DATE", "-", "Y", "保单生效日期", "2026-01-15"),
        ("premium", "DECIMAL", "12,2", "Y", "应缴保费", "1280.00"),
        ("payment_method", "VARCHAR", "8", "Y", "缴费方式：ANNUAL/MONTHLY", "ANNUAL"),
        ("cash_value", "DECIMAL", "12,2", "N", "当前现金价值", "180.00"),
        ("grace_period_end", "DATE", "-", "N", "宽限期截止日期", "2026-07-15"),
        ("loan_balance", "DECIMAL", "12,2", "N", "保单贷款余额", "0.00"),
        ("create_time", "DATETIME", "-", "Y", "创建时间", "2026-01-15 10:30:00"),
        ("update_time", "DATETIME", "-", "Y", "最后更新时间", "2026-05-22 14:00:00"),
    ]
    for row in data:
        ws.append(row)

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=6):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True)

    # 设置列宽
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 8
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 35
    ws.column_dimensions["F"].width = 22

    # Sheet2: 退保操作记录表
    ws2 = wb.create_sheet("退保操作记录表")
    headers2 = ["字段名", "字段类型", "长度", "必填", "说明", "示例值"]
    ws2.append(headers2)
    for col_idx, _ in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    data2 = [
        ("surrender_id", "VARCHAR", "32", "Y", "退保流水号，主键", "SR20260522001"),
        ("policy_no", "VARCHAR", "32", "Y", "保单号", "PL2026050001"),
        ("surrender_type", "VARCHAR", "16", "Y", "退保类型：COOLING_OFF/NORMAL", "COOLING_OFF"),
        ("refund_amount", "DECIMAL", "12,2", "Y", "退费金额", "1270.00"),
        ("cash_value", "DECIMAL", "12,2", "N", "现金价值（正常退保时）", "180.00"),
        ("status", "VARCHAR", "16", "Y", "处理状态：PENDING/COMPLETED/FAILED", "COMPLETED"),
        ("apply_time", "DATETIME", "-", "Y", "申请时间", "2026-05-22 14:30:00"),
        ("complete_time", "DATETIME", "-", "N", "完成时间", "2026-05-22 14:30:05"),
        ("transaction_id", "VARCHAR", "32", "N", "收付交易流水号", "TXN20260522001"),
    ]
    for row in data2:
        ws2.append(row)

    for row in ws2.iter_rows(min_row=2, max_row=ws2.max_row, max_col=6):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True)

    ws2.column_dimensions["A"].width = 18
    ws2.column_dimensions["B"].width = 12
    ws2.column_dimensions["C"].width = 8
    ws2.column_dimensions["D"].width = 8
    ws2.column_dimensions["E"].width = 38
    ws2.column_dimensions["F"].width = 22

    # Sheet3: 现金价值表
    ws3 = wb.create_sheet("现金价值表")
    headers3 = ["产品代码", "保单年度", "投保年龄", "性别", "每万元保额现金价值"]
    ws3.append(headers3)
    for col_idx, _ in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    data3 = [
        ("CRD-CHILD-2026", 1, 0, "M", 18.0),
        ("CRD-CHILD-2026", 2, 0, "M", 52.0),
        ("CRD-CHILD-2026", 3, 0, "M", 98.0),
        ("CRD-CHILD-2026", 5, 0, "M", 235.0),
        ("CRD-CHILD-2026", 10, 0, "M", 680.0),
        ("CRD-CHILD-2026", 15, 0, "M", 1250.0),
        ("CRD-CHILD-2026", 20, 0, "M", 1920.0),
        ("CRD-CHILD-2026", 1, 5, "M", 20.0),
        ("CRD-CHILD-2026", 5, 5, "M", 255.0),
        ("CRD-CHILD-2026", 10, 5, "M", 720.0),
        ("CRD-CHILD-2026", 1, 0, "F", 16.0),
        ("CRD-CHILD-2026", 5, 0, "F", 215.0),
        ("CRD-CHILD-2026", 10, 0, "F", 640.0),
    ]
    for row in data3:
        ws3.append(row)

    for row in ws3.iter_rows(min_row=2, max_row=ws3.max_row, max_col=5):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")

    ws3.column_dimensions["A"].width = 20
    ws3.column_dimensions["B"].width = 12
    ws3.column_dimensions["C"].width = 12
    ws3.column_dimensions["D"].width = 8
    ws3.column_dimensions["E"].width = 22

    path = os.path.join(OUTPUT_DIR, "退保数据字典.xlsx")
    wb.save(path)
    print(f"[OK] 附件A已生成: {path}")


# ============================================================
# 附件B：退保接口规范.docx
# ============================================================
def generate_interface_docx():
    doc = Document()

    title = doc.add_heading("退保接口规范", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("本文档定义网销渠道退保功能涉及的微服务接口，供前端和后端开发参考。")

    # 接口1
    doc.add_heading("1. 退保申请接口", level=1)
    doc.add_heading("1.1 接口信息", level=2)
    doc.add_paragraph("请求方式：POST")
    doc.add_paragraph("请求地址：/api/v1/policy/surrender")
    doc.add_paragraph("Content-Type：application/json")

    doc.add_heading("1.2 请求参数", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["参数名", "类型", "必填", "说明", "示例"]):
        hdr[i].text = h
        set_cell_shading(hdr[i], "4472C4")

    params = [
        ("policyNo", "String", "Y", "保单号", "PL2026050001"),
        ("surrenderType", "String", "Y", "退保类型：COOLING_OFF/NORMAL", "COOLING_OFF"),
        ("applicantId", "String", "Y", "申请人证件号", "310101199001011234"),
        ("bankAccount", "String", "N", "退费收款账户（为空退回原账户）", ""),
        ("reason", "String", "N", "退保原因", "不需要了"),
    ]
    for p in params:
        add_table_row(table, p)

    doc.add_heading("1.3 响应参数", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["参数名", "类型", "说明", "示例"]):
        hdr[i].text = h
        set_cell_shading(hdr[i], "4472C4")

    resp = [
        ("code", "Integer", "状态码，0为成功", "0"),
        ("surrenderAmount", "BigDecimal", "退费金额", "1270.00"),
        ("cashValue", "BigDecimal", "当前现金价值（正常退保时返回）", "180.00"),
        ("processTime", "String", "预计处理时间", "5个工作日"),
        ("message", "String", "提示信息", "退保申请已受理"),
    ]
    for r in resp:
        add_table_row(table, r)

    doc.add_heading("1.4 业务规则", level=2)
    doc.add_paragraph("犹豫期退保：退费金额 = 已缴保费 - 工本费", style="List Bullet")
    doc.add_paragraph("正常退保：退费金额 = 现金价值 - 保单贷款本息", style="List Bullet")
    doc.add_paragraph("退保受理后不可撤回", style="List Bullet")
    doc.add_paragraph("退保成功后调用收付微服务 POST /api/v1/payment/refund 执行退费", style="List Bullet")

    # 接口2
    doc.add_heading("2. 现金价值查询接口", level=1)
    doc.add_heading("2.1 接口信息", level=2)
    doc.add_paragraph("请求方式：GET")
    doc.add_paragraph("请求地址：/api/v1/policy/cash-value")

    doc.add_heading("2.2 请求参数", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["参数名", "类型", "必填", "说明"]):
        hdr[i].text = h
        set_cell_shading(hdr[i], "4472C4")

    params2 = [
        ("policyNo", "String", "Y", "保单号"),
        ("queryDate", "String", "N", "查询日期（默认当天）"),
    ]
    for p in params2:
        add_table_row(table, p)

    doc.add_heading("2.3 响应参数", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["参数名", "类型", "说明", "示例"]):
        hdr[i].text = h
        set_cell_shading(hdr[i], "4472C4")

    resp2 = [
        ("code", "Integer", "状态码", "0"),
        ("policyNo", "String", "保单号", "PL2026050001"),
        ("cashValue", "BigDecimal", "当前现金价值", "180.00"),
        ("policyYear", "Integer", "当前保单年度", "1"),
    ]
    for r in resp2:
        add_table_row(table, r)

    # 错误码
    doc.add_heading("3. 错误码", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "错误码"
    hdr[1].text = "说明"
    set_cell_shading(hdr[0], "4472C4")
    set_cell_shading(hdr[1], "4472C4")

    errors = [
        ("40001", "保单不存在"),
        ("40002", "保单状态不允许退保"),
        ("40003", "申请人身份校验失败"),
        ("40004", "退保金额计算异常"),
        ("40005", "保单贷款未结清，不允许退保"),
        ("50001", "收付微服务调用失败"),
    ]
    for e in errors:
        add_table_row(table, e)

    path = os.path.join(OUTPUT_DIR, "退保接口规范.docx")
    doc.save(path)
    print(f"[OK] 附件B已生成: {path}")


if __name__ == "__main__":
    generate_main_docx()
    generate_data_dict_xlsx()
    generate_interface_docx()
    print("\n全部生成完毕！")
